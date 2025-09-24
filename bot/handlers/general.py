from aiogram import Router, F
from aiogram.types import Message, CallbackQuery, BufferedInputFile, InlineKeyboardMarkup, InlineKeyboardButton
from io import BytesIO
from docx import Document as DocxDocument
from app.core.logging import main_logger
from bot.keyboards.keyboards import get_operator_access_request_keyboard, get_main_keyboard
from bot.service.user_service import UserService
from bot.utils.depend import get_atomic_db
from bot.models.post import PostStatus
from bot.models.user_model import Language, TimeZone
from bot.utils.time_utils import format_dt, get_dt_format
from bot.utils.i18n import t, t_plain, strip_html



router = Router()


def _format_requester_display(user) -> str:
    if getattr(user, "username", None):
        return f"@{user.username}"
    name_parts = [p for p in [getattr(user, 'first_name', None), getattr(user, 'last_name', None)] if p]
    visible_name = " ".join(name_parts) if name_parts else "Пользователь"
    return f"<a href=\"tg://user?id={user.telegram_id}\">{visible_name}</a>"


# ---------- Настройки: helpers ----------

def _settings_main_text(lang: str, tz: str) -> str:
    return (
        f"{t(lang, 'settings_title')}\n\n"
        f"{t(lang, 'settings_lang', lang_code=lang.upper())}\n"
        f"{t(lang, 'settings_tz', tz=tz)}\n\n"
        f"{t(lang, 'settings_main')}"
    )


def _settings_main_keyboard(lang: str) -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=t(lang, 'btn_lang'), callback_data="open_lang")],
        [InlineKeyboardButton(text=t(lang, 'btn_tz'), callback_data="open_tz")],
    ])


def _lang_keyboard(cur_lang: str) -> InlineKeyboardMarkup:
    cur = (cur_lang or Language.RU.value).lower()
    def mark(v: str) -> str:
        return ("✓ " + v.upper()) if v.lower() == cur.lower() else v.upper()
    row = [
        InlineKeyboardButton(text=mark("ru"), callback_data="set_lang:ru"),
        InlineKeyboardButton(text=mark("en"), callback_data="set_lang:en"),
        InlineKeyboardButton(text=mark("es"), callback_data="set_lang:es"),
        InlineKeyboardButton(text=mark("fr"), callback_data="set_lang:fr"),
    ]
    return InlineKeyboardMarkup(inline_keyboard=[row, [InlineKeyboardButton(text=t(cur_lang, 'back'), callback_data="settings_back")]])


def _tz_keyboard(cur_lang: str, cur_tz: str) -> InlineKeyboardMarkup:
    cur_tz = (cur_tz or TimeZone.GMT.value).upper()
    def mark(v: str) -> str:
        return ("✓ " + v) if v.upper() == cur_tz else v
    row1 = [
        InlineKeyboardButton(text=mark("GMT"), callback_data="set_tz:GMT"),
        InlineKeyboardButton(text=mark("UTC"), callback_data="set_tz:UTC"),
        InlineKeyboardButton(text=mark("MSK"), callback_data="set_tz:MSK"),
    ]
    row2 = [
        InlineKeyboardButton(text=mark("CET"), callback_data="set_tz:CET"),
        InlineKeyboardButton(text=mark("EST"), callback_data="set_tz:EST"),
        InlineKeyboardButton(text=mark("PST"), callback_data="set_tz:PST"),
        InlineKeyboardButton(text=mark("IST"), callback_data="set_tz:IST"),
    ]
    row3 = [
        InlineKeyboardButton(text=mark("AEST"), callback_data="set_tz:AEST"),
    ]
    return InlineKeyboardMarkup(inline_keyboard=[row1, row2, row3, [InlineKeyboardButton(text=t(cur_lang, 'back'), callback_data="settings_back")]])


@router.message(F.text.in_({"⚙️ Настройки", "⚙️ Settings"}))
async def show_settings(message: Message):
    try:
        async with get_atomic_db() as db:
            user = await UserService(db).get_user_by_filter(telegram_id=message.from_user.id)
            if not user:
                await message.answer("❌ Пользователь не найден.")
                return
            st = await db.user.get_or_create_settings(user.id)
            lang = st.language or Language.RU.value
            tz = st.time_zone or TimeZone.GMT.value
            await message.answer(_settings_main_text(lang, tz), reply_markup=_settings_main_keyboard(lang))
    except Exception as e:
        main_logger.error(f"show_settings error: {e}")
        await message.answer("⚠️ Ошибка при загрузке настроек.")


@router.callback_query(F.data == "open_lang")
async def open_lang(callback: CallbackQuery):
    try:
        async with get_atomic_db() as db:
            user = await UserService(db).get_user_by_filter(telegram_id=callback.from_user.id)
            if not user:
                await callback.answer("Не найден", show_alert=False)
                return
            st = await db.user.get_or_create_settings(user.id)
            lang = st.language
            await callback.message.edit_text(t(lang, 'choose_lang_title'))
            await callback.message.edit_reply_markup(reply_markup=_lang_keyboard(lang))
        await callback.answer()
    except Exception as e:
        main_logger.error(f"open_lang error: {e}")
        await callback.answer("Ошибка", show_alert=False)


@router.callback_query(F.data == "open_tz")
async def open_tz(callback: CallbackQuery):
    try:
        async with get_atomic_db() as db:
            user = await UserService(db).get_user_by_filter(telegram_id=callback.from_user.id)
            if not user:
                await callback.answer("Не найден", show_alert=False)
                return
            st = await db.user.get_or_create_settings(user.id)
            lang = st.language
            await callback.message.edit_text(t(lang, 'choose_tz_title'))
            await callback.message.edit_reply_markup(reply_markup=_tz_keyboard(lang, st.time_zone))
        await callback.answer()
    except Exception as e:
        main_logger.error(f"open_tz error: {e}")
        await callback.answer("Ошибка", show_alert=False)


@router.callback_query(F.data == "settings_back")
async def settings_back(callback: CallbackQuery):
    try:
        async with get_atomic_db() as db:
            user = await UserService(db).get_user_by_filter(telegram_id=callback.from_user.id)
            if not user:
                await callback.answer("Не найден", show_alert=False)
                return
            st = await db.user.get_or_create_settings(user.id)
            lang = st.language
            tz = st.time_zone
            await callback.message.edit_text(_settings_main_text(lang, tz))
            await callback.message.edit_reply_markup(reply_markup=_settings_main_keyboard(lang))
        await callback.answer()
    except Exception as e:
        main_logger.error(f"settings_back error: {e}")
        await callback.answer("Ошибка", show_alert=False)


@router.callback_query(F.data.startswith("set_lang:"))
async def set_language(callback: CallbackQuery):
    lang_value = callback.data.split(":", 1)[1].lower()
    if lang_value not in {x.value for x in Language}:
        await callback.answer("Недопустимый язык", show_alert=False)
        return
    try:
        async with get_atomic_db() as db:
            user = await UserService(db).get_user_by_filter(telegram_id=callback.from_user.id)
            if not user:
                await callback.answer("Не найден", show_alert=False)
                return
            await db.user.update_settings(user.id, {"language": lang_value})
            st = await db.user.get_settings(user.id)
            cur_lang = st.language
            # Обновляем заголовок и клавиатуру экрана выбора языка
            await callback.message.edit_text(t(cur_lang, 'choose_lang_title'))
            await callback.message.edit_reply_markup(reply_markup=_lang_keyboard(cur_lang))
        await callback.answer(t(lang_value, 'saved'))
        # Обновляем главное меню (ReplyKeyboard) в соответствии с новым языком
        # Определяем роль
        is_admin = getattr(user, 'role', '') == 'admin' or getattr(user, 'is_admin', False)
        is_operator = getattr(user, 'role', '') == 'operator' or getattr(user, 'is_operator', False)
        try:
            await callback.message.answer(
                t(lang_value, 'saved'),
                reply_markup=get_main_keyboard(lang_value, is_admin=is_admin, is_operator=is_operator)
            )
        except Exception:
            pass
    except Exception as e:
        main_logger.error(f"set_language error: {e}")
        await callback.answer("Ошибка", show_alert=False)


@router.callback_query(F.data.startswith("set_tz:"))
async def set_time_zone(callback: CallbackQuery):
    tz_value = callback.data.split(":", 1)[1].upper()
    if tz_value not in {x.value for x in TimeZone}:
        await callback.answer("Недопустимая TZ", show_alert=False)
        return
    try:
        async with get_atomic_db() as db:
            user = await UserService(db).get_user_by_filter(telegram_id=callback.from_user.id)
            if not user:
                await callback.answer("Не найден", show_alert=False)
                return
            await db.user.update_settings(user.id, {"time_zone": tz_value})
            st = await db.user.get_settings(user.id)
            cur_lang = st.language
            # Обновляем заголовок и клавиатуру экрана выбора TZ
            await callback.message.edit_text(t(cur_lang, 'choose_tz_title'))
            await callback.message.edit_reply_markup(reply_markup=_tz_keyboard(cur_lang, st.time_zone))
        await callback.answer(t(cur_lang, 'saved'))
    except Exception as e:
        main_logger.error(f"set_time_zone error: {e}")
        await callback.answer("Ошибка", show_alert=False)


@router.message(F.text.in_(
    {"📊 Отчет", "📊 Отчёт", "Отчёт", "Отчет", "📊 Report"}
))
async def show_report(message: Message):
    within_hours = 24
    try:
        async with get_atomic_db() as db:
            # Получаем язык/часовой пояс вызывающего
            user = await UserService(db).get_user_by_filter(telegram_id=message.from_user.id)
            st = await db.user.get_or_create_settings(user.id) if user else None
            lang = st.language if st else Language.RU.value
            tz = st.time_zone if st else TimeZone.GMT.value
            fmt = get_dt_format(lang)

            total_channels = await db.channel.count_channels()
            total_keywords = len(await db.keywords.get_all_keywords())
            total_matched_posts = await db.post.count_distinct_posts_with_matches(within_hours)
            processed = await db.post.count_processing_by_status(PostStatus.PROCESSED.value, within_hours)
            postponed = await db.post.count_processing_by_status(PostStatus.POSTPONED.value, within_hours)
            pending = await db.post.count_processing_by_status(PostStatus.PENDING.value, within_hours)

            # Статистика по операторам (оставим формат строк как есть)
            op_stats_raw = await db.post.get_operator_stats(within_hours)
            op_lines = []
            for op_id, proc_cnt, postp_cnt in op_stats_raw:
                u = await db.user.get_user_by_filter(id=op_id)
                if u and getattr(u, "username", None):
                    display = f"@{u.username}"
                elif u:
                    name_parts = [p for p in [getattr(u, 'first_name', None), getattr(u, 'last_name', None)] if p]
                    visible = " ".join(name_parts) if name_parts else str(u.telegram_id)
                    display = f"<a href=\"tg://user?id={u.telegram_id}\">{visible}</a>"
                else:
                    display = "—"
                # пока без локализации подписи
                op_lines.append(f"• {display}: обработано - <b>{proc_cnt}</b>, отложено - <b>{postp_cnt}</b>")

        text = (
            f"{t(lang, 'report_title', hours=within_hours)}\n\n"
            f"{t(lang, 'total_channels', n=total_channels)}\n"
            f"{t(lang, 'total_keywords', n=total_keywords)}\n"
            f"{t(lang, 'found_posts', n=total_matched_posts)}\n"
            f"{t(lang, 'processed', n=processed)}\n"
            f"{t(lang, 'postponed', n=postponed)}\n"
            f"{t(lang, 'pending', n=pending)}"
            + (t(lang, 'operators_header') + "\n" + ("\n".join(op_lines) if op_lines else t(lang, 'dash')))
        )
        await message.answer(text, disable_web_page_preview=True)

        # DOCX отчёт
        if DocxDocument is None:
            await message.answer("⚠️ Подробный .docx отчёт недоступен (python-docx не установлен).")
            return

        try:
            async with get_atomic_db() as db:
                posts = await db.post.get_recent_matched_posts(within_hours)

            doc = DocxDocument()
            doc.add_heading(t_plain(lang, 'report_title', hours=within_hours), level=0)
            doc.add_paragraph("")
            doc.add_paragraph(t_plain(lang, 'total_channels', n=total_channels))
            doc.add_paragraph(t_plain(lang, 'total_keywords', n=total_keywords))
            doc.add_paragraph(t_plain(lang, 'found_posts', n=total_matched_posts))
            doc.add_paragraph(t_plain(lang, 'processed', n=processed))
            doc.add_paragraph(t_plain(lang, 'postponed', n=postponed))
            doc.add_paragraph(t_plain(lang, 'pending', n=pending))

            doc.add_heading("Операторы", level=1)
            if op_lines:
                for line in op_lines:
                    doc.add_paragraph(strip_html(line), style="List Bullet")
            else:
                doc.add_paragraph("—")

            doc.add_heading("Посты", level=1)
            for p in posts:
                ch_title = p.channel.title if p.channel.title else p.channel.channel_username
                doc.add_heading(ch_title, level=2)
                doc.add_paragraph(t_plain(lang, 'notify_date', dt=format_dt(getattr(p, 'published_at', None), tz, fmt)))
                if getattr(p, "url", None):
                    doc.add_paragraph(f"URL: {p.url}")
                kw_texts = []
                try:
                    for mk in getattr(p, "matched_keywords", []) or []:
                        if getattr(mk, "keyword", None) and mk.keyword.text:
                            if mk.keyword.text not in kw_texts:
                                kw_texts.append(mk.keyword.text)
                except Exception:
                    pass
                if kw_texts:
                    doc.add_paragraph("Keywords: " + ", ".join(kw_texts))
                preview = (p.text or "").strip()
                doc.add_paragraph(strip_html(preview) if preview else "(no text)")
                doc.add_paragraph("")

            bio = BytesIO()
            doc.save(bio)
            bio.seek(0)
            fname = f"report_{within_hours}h.docx"
            file = BufferedInputFile(bio.read(), filename=fname)
            await message.answer_document(file, caption=t(lang, 'detailed_report_caption'))
        except Exception as e:
            main_logger.error(f"show_report docx error: {e}")
            await message.answer("⚠️ Не удалось сформировать отчёт. Попробуйте позже.")

    except Exception as e:
        main_logger.error(f"show_report error: {e}")
        await message.answer("⚠️ Не удалось сформировать отчёт. Попробуйте позже.")


@router.message(F.text == "📝 Получить доступ оператора")
async def request_operator_access(message: Message):
    """Обычный пользователь запрашивает роль оператора: уведомляем всех админов."""
    try:
        async with get_atomic_db() as db:
            user_service = UserService(db)
            user = await user_service.get_user_by_filter(telegram_id=message.from_user.id)
            if not user:
                await message.answer("❌ Не удалось определить пользователя. Попробуйте /start.")
                return

            admins = await user_service.get_admins()
            if not admins:
                await message.answer("⚠️ Нет доступных администраторов. Попробуйте позже.")
                return

            display = _format_requester_display(user)
            text = (
                "📝 <b>Запрос на доступ оператора</b>\n\n"
                f"Запрос от: {display}\n"
                "Нажмите, чтобы выдать доступ или отклонить."
            )
            kb = get_operator_access_request_keyboard(user.id)

            sent = 0
            for admin in admins:
                try:
                    await message.bot.send_message(admin.telegram_id, text, reply_markup=kb)
                    sent += 1
                except Exception as e:
                    main_logger.error(f"send admin op-access request error to {admin.telegram_id}: {e}")

        if sent:
            await message.answer(
                "✅ <b>Запрос отправлен!</b>\n\n"
                "Администраторы получили ваше обращение.\n"
                "Вы получите уведомление после принятия решения."
            )
        else:
            await message.answer(
                "❌ <b>Не удалось отправить запрос</b>\n\n"
                "Попробуйте позже или свяжитесь с поддержкой."
            )
    except Exception as e:
        main_logger.error(f"request_operator_access error: {e}")
        await message.answer("⚠️ Произошла ошибка. Попробуйте позже.")


@router.callback_query(F.data.startswith("approve_operator:"))
async def approve_operator(callback: CallbackQuery):
    # Проверка прав администратора
    async with get_atomic_db() as db:
        perms = await UserService(db).cheek_user_permissions(callback.from_user.id)
        if not perms.get("is_admin"):
            await callback.answer("Нет прав", show_alert=False)
            return

    try:
        user_id = int(callback.data.split(":")[1])
    except Exception:
        await callback.answer("Некорректные данные", show_alert=False)
        return

    try:
        async with get_atomic_db() as db:
            svc = UserService(db)
            user = await svc.get_user_by_filter(id=user_id)
            if not user:
                await callback.answer("Пользователь не найден", show_alert=False)
                return
            updated = await svc.set_operator(user_id, True)
            # Уведомляем целевого пользователя
            try:
                await callback.bot.send_message(
                    chat_id=user.telegram_id,
                    text=(
                        "✅ <b>Доступ оператора выдан</b>\n\n"
                        "Перезапустите бота командой /start, чтобы обновить меню."
                    )
                )
            except Exception as e:
                main_logger.error(f"notify user grant operator failed {user.telegram_id}: {e}")
    except Exception as e:
        main_logger.error(f"approve_operator error: {e}")
        await callback.answer("Ошибка", show_alert=False)
        return

    try:
        await callback.message.edit_text("✅ Запрос обработан: доступ выдан.")
    except Exception:
        pass
    await callback.answer()


@router.callback_query(F.data.startswith("reject_operator:"))
async def reject_operator(callback: CallbackQuery):
    # Проверка прав администратора
    async with get_atomic_db() as db:
        perms = await UserService(db).cheek_user_permissions(callback.from_user.id)
        if not perms.get("is_admin"):
            await callback.answer("Нет прав", show_alert=False)
            return

    try:
        user_id = int(callback.data.split(":")[1])
    except Exception:
        await callback.answer("Некорректные данные", show_alert=False)
        return

    try:
        async with get_atomic_db() as db:
            svc = UserService(db)
            user = await svc.get_user_by_filter(id=user_id)
            if not user:
                await callback.answer("Пользователь не найден", show_alert=False)
                return
            # Можно уведомить пользователя
            try:
                await callback.bot.send_message(
                    chat_id=user.telegram_id,
                    text="❌ Ваш запрос на доступ оператора отклонён."
                )
            except Exception as e:
                main_logger.error(f"notify user reject operator failed {user.telegram_id}: {e}")
    except Exception as e:
        main_logger.error(f"reject_operator error: {e}")
        await callback.answer("Ошибка", show_alert=False)
        return

    try:
        await callback.message.edit_text("❌ Запрос отклонён.")
    except Exception:
        pass
    await callback.answer()


@router.message(F.text == "💬 Обратная связь")
async def show_feedback(message: Message):
    text = (
        "💬 <b>Обратная связь</b>\n\n"
        "Связь с поддержкой:\n"
        "• Telegram: <a href=\"https://t.me/support\">@support</a>\n"
        "• Email: <code>support@example.com</code>\n\n"
        "🕘 График: пн–пт, 10:00–19:00 (МСК)"
    )
    await message.answer(text)


@router.message(F.text.in_({"❓О системе", "❓ О системе", "❓About", "❓ About", "❓About", "❓About"}))
async def about_system(message: Message):
    text = (
        "❓ <b>О системе</b>\n\n"
        "Система мониторинга Telegram‑каналов помогает команде оперативно работать с контентом:\n\n"
        "• 📢 Добавление и модерация каналов.\n"
        "• 🔍 Поиск по ключевым словам (предложения операторов, подтверждение админами).\n"
        "• 👤 Роли и права: администратор и оператор.\n"
        "• 🔔 Уведомления и удобные панели управления прямо в чате.\n\n"
        "💡 Чтобы начать — воспользуйтесь кнопками меню ниже."
    )
    await message.answer(text)
